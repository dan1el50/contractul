"""The phase 1 rendering spike, as tests.

The question phase 1 exists to answer: can headless LibreOffice convert our
.docx templates to PDF without mangling Romanian diacritics or Cyrillic?

Font substitution is the specific risk. LibreOffice silently swaps a missing
font for whatever it has, and a missing glyph becomes a box in the PDF with
no error raised anywhere — the conversion "succeeds" and the document is
ruined. So these tests read the text back out of the produced PDF and assert
the characters survived. Asserting the PDF merely exists would pass on
exactly the failure we are worried about.

The fixture is a real bilingual RO/RU document with embedded screenshots. It
is a TEST FIXTURE, not catalog content — it is an internal 2FA guide, not a
contract. Real templates come from the legal team.
"""

from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader

from app.documents.renderer import (
    RenderError,
    docx_to_pdf,
    fill_template,
    pdf_to_previews,
)

FIXTURE = Path(__file__).parent.parent / "fixtures" / "activare-2fa.docx"

LATIN_SAMPLES = ["Activare 2FA", "Urmariti instructiunile"]
CYRILLIC_SAMPLES = ["Активация 2FA", "Следуйте инструкциям"]

# The fixture is typed WITHOUT diacritics (informal Moldovan style — "Urmariti",
# not "Urmăriți") and contains no placeholders. So on its own it proves neither
# that diacritics survive nor that template filling works, and a test suite built
# only on it would report a green phase 1 while leaving both open.
#
# Hence the synthetic template below: it carries every Romanian diacritic, both
# scripts in one file, and real placeholders. Generated rather than committed so
# that what it covers is readable here rather than hidden in a binary.
ROMANIAN_DIACRITICS = "ăâîșțĂÂÎȘȚ"

# Romanian's ș/ț are comma-below (U+0219/U+021B), a different codepoint from the
# cedilla forms ş/ţ (U+015F/U+0163). Substitution silently swaps one for the
# other, which is wrong but not visibly broken — worth asserting exactly.
COMMA_BELOW = "șț"
CEDILLA = "şţ"


@pytest.fixture(scope="module")
def rendered_pdf(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Fill and convert the fixture once; the assertions below all read it."""
    workdir = tmp_path_factory.mktemp("render")
    filled = fill_template(FIXTURE, {}, workdir / "filled.docx")
    return docx_to_pdf(filled, workdir / "out")


def _pdf_text(pdf_path: Path) -> str:
    return "\n".join(page.extract_text() for page in PdfReader(str(pdf_path)).pages)


def test_fixture_exists() -> None:
    assert FIXTURE.is_file(), f"Missing test fixture: {FIXTURE}"


def test_template_without_placeholders_passes_through(tmp_path: Path) -> None:
    """A fixed-text document is a valid template, not an error."""
    output = fill_template(FIXTURE, {}, tmp_path / "out.docx")

    assert output.is_file()
    assert output.stat().st_size > 0


def test_missing_template_raises_render_error(tmp_path: Path) -> None:
    with pytest.raises(RenderError, match="Template not found"):
        fill_template(tmp_path / "nope.docx", {}, tmp_path / "out.docx")


def test_conversion_produces_a_pdf(rendered_pdf: Path) -> None:
    assert rendered_pdf.is_file()
    assert rendered_pdf.read_bytes().startswith(b"%PDF")


def test_page_count_survives_conversion(rendered_pdf: Path) -> None:
    # The source document is 2 pages. A collapsed or exploded page count means
    # layout was not preserved.
    assert len(PdfReader(str(rendered_pdf)).pages) == 2


@pytest.mark.parametrize("sample", LATIN_SAMPLES)
def test_latin_text_survives_conversion(rendered_pdf: Path, sample: str) -> None:
    """Note: the fixture has no diacritics, so this covers plain Latin only.

    Diacritics are covered by the synthetic template further down.
    """
    assert sample in _pdf_text(rendered_pdf)


@pytest.mark.parametrize("sample", CYRILLIC_SAMPLES)
def test_cyrillic_text_survives_conversion(rendered_pdf: Path, sample: str) -> None:
    """The load-bearing assertion of the whole phase.

    Romanian and Russian coexist in one file, so a font that covers only Latin
    silently destroys half of every document we sell.
    """
    assert sample in _pdf_text(rendered_pdf)


def test_no_replacement_characters_in_output(rendered_pdf: Path) -> None:
    """Catches the substitution failure directly.

    A glyph LibreOffice could not render becomes U+FFFD or a box rather than
    an exception. Nothing else in the pipeline would notice.
    """
    text = _pdf_text(rendered_pdf)

    assert "�" not in text, "Replacement character in PDF — a font is missing glyphs"


def test_previews_are_generated_one_per_page(rendered_pdf: Path, tmp_path: Path) -> None:
    pages = pdf_to_previews(rendered_pdf, tmp_path / "previews")

    assert len(pages) == 2
    assert all(page.read_bytes().startswith(b"\x89PNG") for page in pages)


# ─── Synthetic template: diacritics + placeholders ────────────────────────────
#
# What the real fixture cannot tell us, because it has neither.


@pytest.fixture(scope="module")
def synthetic_pdf(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """A template with both scripts and real placeholders, filled and converted."""
    workdir = tmp_path_factory.mktemp("synthetic")
    template_path = workdir / "template.docx"

    document = Document()
    document.add_paragraph("CONTRACT DE PRESTĂRI SERVICII")
    document.add_paragraph(f"Toate diacriticele: {ROMANIAN_DIACRITICS}")
    document.add_paragraph("Părțile au convenit următoarele condiții:")
    document.add_paragraph("Beneficiar: {{ client_name }}")
    document.add_paragraph("Suma: {{ amount }} MDL")
    # Both scripts in one file — the model the legal team confirmed.
    document.add_paragraph("Стороны договорились о следующем:")
    document.add_paragraph("Заказчик: {{ client_name }}")
    document.save(str(template_path))

    filled = fill_template(
        template_path,
        {"client_name": 'SRL "NordConstruct"', "amount": "900"},
        workdir / "filled.docx",
    )
    return docx_to_pdf(filled, workdir / "out")


def test_placeholders_are_filled(synthetic_pdf: Path) -> None:
    text = _pdf_text(synthetic_pdf)

    assert 'SRL "NordConstruct"' in text
    assert "900 MDL" in text
    # An unrendered placeholder means docxtpl silently did nothing.
    assert "{{" not in text


def test_every_romanian_diacritic_survives_conversion(synthetic_pdf: Path) -> None:
    """The assertion phase 1 exists for.

    A font missing these renders boxes, and nothing raises.
    """
    text = _pdf_text(synthetic_pdf)

    missing = [ch for ch in ROMANIAN_DIACRITICS if ch not in text]
    assert not missing, f"Diacritics lost in conversion: {''.join(missing)}"


def test_comma_below_is_not_swapped_for_cedilla(synthetic_pdf: Path) -> None:
    """Romanian ș/ț must stay comma-below, not become Turkish cedilla ş/ţ.

    Both render as a recognisable letter, so this failure survives a visual
    check and only shows up when someone copies text out of the document.
    """
    text = _pdf_text(synthetic_pdf)

    assert all(ch in text for ch in COMMA_BELOW)
    assert not any(ch in text for ch in CEDILLA), "Comma-below substituted with cedilla"


def test_both_scripts_coexist_in_one_document(synthetic_pdf: Path) -> None:
    """One file carries every language — so a Latin-only font ruins half of it."""
    text = _pdf_text(synthetic_pdf)

    assert "Părțile au convenit" in text
    assert "Стороны договорились" in text
    assert "�" not in text
